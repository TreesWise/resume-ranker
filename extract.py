import fitz  # PyMuPDF
from docx import Document
import io

def extract_text_from_pdf_or_docx(file_content: bytes, filename: str) -> str:
    """
    Unified function to extract text from PDF or DOCX based on file extension.
    """
    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_content)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(file_content)
    else:
        return ""

def extract_text_from_pdf(file_content: bytes) -> str:
    doc = fitz.open(stream=file_content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

from docx import Document
import io

# def extract_text_from_docx(file_like):
#     try:
#         doc = Document(io.BytesIO(file_like.read()))
#         text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        
#         # Extract text from tables
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip() != ""])
#                 text += "\n" + row_text
        
#         # Extract text from headers (heading levels 1-9)
#         for paragraph in doc.paragraphs:
#             if paragraph.style.name.startswith('Heading'):
#                 text += "\n" + paragraph.text.strip()

#         return text

#     except Exception as e:
#         print(f"[ERROR] Failed to extract from docx: {e}")
#         return None
def extract_text_from_docx(file_like):
    try:
        doc = Document(io.BytesIO(file_like.read()))
        seen_rows = set()
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                row_text_clean = row_text.lower().strip()
                if row_text_clean and row_text_clean not in seen_rows:
                    text += "\n" + row_text
                    seen_rows.add(row_text_clean)

        # Extract text from headers (optional — can already be in paragraphs)
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                header_text = paragraph.text.strip()
                if header_text and header_text.lower() not in seen_rows:
                    text += "\n" + header_text
                    seen_rows.add(header_text.lower())
        text = remove_duplicate_lines(text)
        print(text)
        return text

    except Exception as e:
        print(f"[ERROR] Failed to extract from docx: {e}")
        return None

def remove_duplicate_lines(text: str) -> str:
    lines = text.split("\n")
    seen = set()
    cleaned = []
    for line in lines:
        key = line.strip().lower()
        if key and key not in seen:
            cleaned.append(line)
            seen.add(key)
    return "\n".join(cleaned)
